"""DocumentProcessorAgent — read, extract, summarize, and create PDF/Word/Excel files.

Library availability is checked at import time with graceful degradation:
  - PDF read   : pypdf  (pip install pypdf)
  - PDF create : fpdf2  (pip install fpdf2)
  - Word       : python-docx (pip install python-docx)
  - Excel      : openpyxl    (pip install openpyxl)

All four libraries are optional; each capability reports an informative error
when the backing library is not installed rather than crashing the platform.

Supported operations
--------------------
  read_pdf        — extract all text from a local PDF file
  read_word       — extract all text + basic structure from a .docx file
  read_excel      — read sheets/rows from a .xlsx/.xls file
  create_pdf      — generate a simple text-based PDF file
  create_word     — generate a .docx file with headings and paragraphs
  create_excel    — generate a .xlsx file from a 2-D data list
  summarize_document — read any supported file and summarise via LLM
"""

from __future__ import annotations

import asyncio
import json
import os
from pathlib import Path
from typing import Any, ClassVar

import structlog

from vclaw.agents.base import AgentBase
from vclaw.domain.models import (
    AgentCapability,
    AgentManifest,
    AgentRequest,
    AgentResponse,
    LLMRequest,
    RetryPolicy,
    ToolDefinition,
)

logger: structlog.stdlib.BoundLogger = structlog.get_logger(__name__)

# ---------------------------------------------------------------------------
# Optional-dependency feature flags
# ---------------------------------------------------------------------------

try:
    import pypdf  # noqa: F401

    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False
    logger.info("pypdf_not_installed", hint="pip install pypdf")

try:
    from fpdf import FPDF  # noqa: F401

    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False
    logger.info("fpdf2_not_installed", hint="pip install fpdf2")

try:
    import docx  # python-docx  # noqa: F401

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.info("python_docx_not_installed", hint="pip install python-docx")

try:
    import openpyxl  # noqa: F401

    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False
    logger.info("openpyxl_not_installed", hint="pip install openpyxl")

# Maximum characters extracted from a single file before truncation
_MAX_EXTRACT_CHARS = 20_000
# Default output directory for generated files
_DEFAULT_OUTPUT_DIR = "/tmp/vclaw_documents"


def _ensure_output_dir(directory: str = _DEFAULT_OUTPUT_DIR) -> Path:
    path = Path(directory)
    path.mkdir(parents=True, exist_ok=True)
    return path


class DocumentProcessorAgent(AgentBase):
    """Agent that reads and creates PDF, Word (.docx), and Excel (.xlsx) files.

    All file-system operations are sandboxed under the configured output directory
    for write operations. Read operations accept any absolute path provided in the
    request input.

    The `summarize_document` tool chains a read operation with an LLM call so the
    caller gets a human-readable summary without having to handle raw text themselves.
    """

    manifest: ClassVar[AgentManifest] = AgentManifest(
        name="document_processor",
        version="0.1.0",
        description=(
            "Document processing agent: reads text from PDF, Word (.docx), and Excel (.xlsx) files; "
            "creates simple documents; and summarises file contents via LLM. "
            "Handles Vietnamese and multilingual documents."
        ),
        capabilities=[
            AgentCapability(
                name="document_reading",
                description="Extract text and structured data from PDF, Word, and Excel files",
            ),
            AgentCapability(
                name="document_creation",
                description="Generate PDF, Word (.docx), and Excel (.xlsx) files with provided content",
            ),
            AgentCapability(
                name="document_summarization",
                description="Read a document file and produce an LLM-generated summary",
            ),
        ],
        tools=[
            ToolDefinition(
                name="read_pdf",
                description="Extract all text from a PDF file. Returns page-by-page text.",
                parameters={
                    "file_path": {"type": "string", "description": "Absolute path to the PDF file"},
                    "max_pages": {
                        "type": "integer",
                        "description": "Maximum number of pages to read (default: all pages)",
                    },
                },
                required_params=["file_path"],
            ),
            ToolDefinition(
                name="read_word",
                description="Extract text and structure from a Word (.docx) file.",
                parameters={
                    "file_path": {"type": "string", "description": "Absolute path to the .docx file"},
                },
                required_params=["file_path"],
            ),
            ToolDefinition(
                name="read_excel",
                description="Read rows and sheets from an Excel (.xlsx or .xls) file.",
                parameters={
                    "file_path": {"type": "string", "description": "Absolute path to the Excel file"},
                    "sheet_name": {
                        "type": "string",
                        "description": "Specific sheet name to read (default: first sheet)",
                    },
                    "max_rows": {
                        "type": "integer",
                        "description": "Maximum number of rows to return per sheet (default: 100)",
                    },
                },
                required_params=["file_path"],
            ),
            ToolDefinition(
                name="create_pdf",
                description="Create a PDF file from plain text content.",
                parameters={
                    "filename": {"type": "string", "description": "Output filename (e.g. report.pdf)"},
                    "title": {"type": "string", "description": "Document title shown at the top"},
                    "content": {"type": "string", "description": "Body text to write into the PDF"},
                    "output_dir": {
                        "type": "string",
                        "description": f"Output directory (default: {_DEFAULT_OUTPUT_DIR})",
                    },
                },
                required_params=["filename", "content"],
            ),
            ToolDefinition(
                name="create_word",
                description="Create a Word (.docx) file with a title, optional headings, and paragraphs.",
                parameters={
                    "filename": {"type": "string", "description": "Output filename (e.g. report.docx)"},
                    "title": {"type": "string", "description": "Document title (Heading 1)"},
                    "sections": {
                        "type": "array",
                        "description": (
                            'List of sections: [{"heading": "Section 1", "body": "text..."}, ...]'
                            " — heading is optional"
                        ),
                    },
                    "output_dir": {
                        "type": "string",
                        "description": f"Output directory (default: {_DEFAULT_OUTPUT_DIR})",
                    },
                },
                required_params=["filename", "sections"],
            ),
            ToolDefinition(
                name="create_excel",
                description="Create an Excel (.xlsx) file from a list of rows.",
                parameters={
                    "filename": {"type": "string", "description": "Output filename (e.g. data.xlsx)"},
                    "sheet_name": {"type": "string", "description": "Sheet name (default: Sheet1)"},
                    "headers": {
                        "type": "array",
                        "description": "Column header names e.g. ['Name', 'Age', 'Score']",
                    },
                    "rows": {
                        "type": "array",
                        "description": "2-D list of row data e.g. [['Alice', 30, 95], ['Bob', 25, 88]]",
                    },
                    "output_dir": {
                        "type": "string",
                        "description": f"Output directory (default: {_DEFAULT_OUTPUT_DIR})",
                    },
                },
                required_params=["filename", "rows"],
            ),
            ToolDefinition(
                name="summarize_document",
                description=(
                    "Read a PDF, Word, or Excel file and generate a concise LLM summary. "
                    "Supports any language including Vietnamese."
                ),
                parameters={
                    "file_path": {"type": "string", "description": "Absolute path to the document"},
                    "language": {
                        "type": "string",
                        "description": "Language for the summary response (e.g. 'Vietnamese', 'English'). Default: auto-detect.",
                    },
                    "focus": {
                        "type": "string",
                        "description": "Optional focus area for the summary (e.g. 'key decisions', 'financial figures')",
                    },
                },
                required_params=["file_path"],
            ),
        ],
        max_concurrent=4,
        timeout_seconds=120.0,
        retry_policy=RetryPolicy(max_retries=1, base_delay_seconds=2.0),
        tags=["document", "pdf", "word", "excel", "file", "summarize"],
    )

    # ------------------------------------------------------------------
    # Main entry point
    # ------------------------------------------------------------------

    async def execute(self, request: AgentRequest) -> AgentResponse:
        text = request.input_data.get("text", "")
        tool_name: str = request.input_data.get("tool", "")
        tool_args: dict[str, Any] = request.input_data.get("args", {})

        # If the caller already specified a tool + args, dispatch directly
        if tool_name and tool_args:
            result = await self._execute_tool(tool_name, tool_args)
            success = result.get("success", False)
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=success,
                data=result.get("data", {}),
                error=result.get("error") if not success else None,
            )

        if not text:
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=False,
                error="No input text or tool/args provided",
            )

        # Route via LLM tool-calling
        try:
            llm_resp = await self.call_llm(
                LLMRequest(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a document processing assistant. "
                                "Use the available tools to read, create, or summarise documents "
                                "as requested by the user. Always select and call the most appropriate tool.\n\n"
                                f"Available libraries — pypdf: {_PYPDF_AVAILABLE}, "
                                f"fpdf2: {_FPDF_AVAILABLE}, "
                                f"python-docx: {_DOCX_AVAILABLE}, "
                                f"openpyxl: {_OPENPYXL_AVAILABLE}"
                            ),
                        },
                        {"role": "user", "content": text},
                    ],
                    tools=self.get_tool_schemas(),
                    tool_choice="auto",
                    temperature=0.0,
                )
            )
        except Exception as exc:
            return self._no_llm_response(request, str(exc))

        if llm_resp.tool_calls:
            return await self._dispatch_tool_calls(request, llm_resp.tool_calls)

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=True,
            data={"response_text": llm_resp.content or "Document task completed."},
        )

    # ------------------------------------------------------------------
    # Tool dispatch helpers
    # ------------------------------------------------------------------

    async def _dispatch_tool_calls(
        self, request: AgentRequest, tool_calls: list[dict[str, Any]]
    ) -> AgentResponse:
        tasks = []
        for tc in tool_calls:
            func = tc.get("function", {})
            name = func.get("name", "")
            try:
                args: dict[str, Any] = json.loads(func.get("arguments", "{}"))
            except json.JSONDecodeError:
                args = {}
            tasks.append(self._execute_tool(name, args))

        results = await asyncio.gather(*tasks, return_exceptions=True)

        all_data: list[dict[str, Any]] = []
        response_parts: list[str] = []
        overall_success = True

        for tc, result in zip(tool_calls, results, strict=False):
            func = tc.get("function", {})
            name = func.get("name", "")
            if isinstance(result, Exception):
                overall_success = False
                response_parts.append(f"Tool {name} error: {result}")
                continue
            assert isinstance(result, dict)
            success = result.get("success", False)
            if not success:
                overall_success = False
                response_parts.append(f"Tool {name} failed: {result.get('error', '?')}")
            else:
                all_data.append({"tool": name, "data": result.get("data", {})})
                response_parts.append(self._format_result(name, result.get("data", {})))

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=overall_success,
            data={
                "response_text": "\n\n".join(response_parts),
                "tool_results": all_data,
            },
        )

    def _format_result(self, tool_name: str, data: dict[str, Any]) -> str:
        if tool_name == "read_pdf":
            pages = data.get("pages", 0)
            preview = data.get("text", "")[:400]
            return f"📄 PDF read — {pages} page(s)\n{preview}..."
        if tool_name == "read_word":
            paragraphs = data.get("paragraph_count", 0)
            preview = data.get("text", "")[:400]
            return f"📝 Word doc read — {paragraphs} paragraph(s)\n{preview}..."
        if tool_name == "read_excel":
            sheets = data.get("sheets_read", [])
            row_count = data.get("total_rows", 0)
            return f"📊 Excel read — sheets: {sheets}, {row_count} row(s)"
        if tool_name in ("create_pdf", "create_word", "create_excel"):
            path = data.get("file_path", "")
            size = data.get("size_bytes", "?")
            return f"✅ Created `{path}` ({size} bytes)"
        if tool_name == "summarize_document":
            summary = data.get("summary", "")
            return f"📋 Summary:\n{summary}"
        return str(data)

    # ------------------------------------------------------------------
    # Tool implementations
    # ------------------------------------------------------------------

    async def _execute_tool(self, name: str, args: dict[str, Any]) -> dict[str, Any]:
        try:
            if name == "read_pdf":
                return await asyncio.get_event_loop().run_in_executor(None, self._read_pdf, args)
            if name == "read_word":
                return await asyncio.get_event_loop().run_in_executor(None, self._read_word, args)
            if name == "read_excel":
                return await asyncio.get_event_loop().run_in_executor(None, self._read_excel, args)
            if name == "create_pdf":
                return await asyncio.get_event_loop().run_in_executor(None, self._create_pdf, args)
            if name == "create_word":
                return await asyncio.get_event_loop().run_in_executor(None, self._create_word, args)
            if name == "create_excel":
                return await asyncio.get_event_loop().run_in_executor(None, self._create_excel, args)
            if name == "summarize_document":
                return await self._summarize_document(args)
            return {"success": False, "error": f"Unknown tool: {name}"}
        except Exception as exc:
            logger.exception("document_tool_error", tool=name, args=args)
            return {"success": False, "error": str(exc)}

    # --- Read tools (synchronous, executed in thread pool) ---

    def _read_pdf(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _PYPDF_AVAILABLE:
            return {"success": False, "error": "pypdf not installed. Run: pip install pypdf"}

        import pypdf as _pypdf

        file_path: str = args["file_path"]
        max_pages: int | None = args.get("max_pages")

        if not os.path.isfile(file_path):
            return {"success": False, "error": f"File not found: {file_path}"}

        reader = _pypdf.PdfReader(file_path)
        total_pages = len(reader.pages)
        pages_to_read = reader.pages[:max_pages] if max_pages else reader.pages

        extracted: list[str] = []
        for i, page in enumerate(pages_to_read):
            page_text = page.extract_text() or ""
            extracted.append(f"--- Page {i + 1} ---\n{page_text}")

        full_text = "\n".join(extracted)
        truncated = len(full_text) > _MAX_EXTRACT_CHARS
        return {
            "success": True,
            "data": {
                "file_path": file_path,
                "total_pages": total_pages,
                "pages": len(extracted),
                "text": full_text[:_MAX_EXTRACT_CHARS],
                "truncated": truncated,
            },
        }

    def _read_word(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        import docx as _docx

        file_path: str = args["file_path"]
        if not os.path.isfile(file_path):
            return {"success": False, "error": f"File not found: {file_path}"}

        doc = _docx.Document(file_path)
        paragraphs: list[dict[str, str]] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({"style": para.style.name, "text": para.text})

        full_text = "\n".join(p["text"] for p in paragraphs)
        truncated = len(full_text) > _MAX_EXTRACT_CHARS
        return {
            "success": True,
            "data": {
                "file_path": file_path,
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs[:50],
                "text": full_text[:_MAX_EXTRACT_CHARS],
                "truncated": truncated,
            },
        }

    def _read_excel(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _OPENPYXL_AVAILABLE:
            return {"success": False, "error": "openpyxl not installed. Run: pip install openpyxl"}

        import openpyxl as _openpyxl

        file_path: str = args["file_path"]
        sheet_name: str | None = args.get("sheet_name")
        max_rows: int = int(args.get("max_rows", 100))

        if not os.path.isfile(file_path):
            return {"success": False, "error": f"File not found: {file_path}"}

        wb = _openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets_data: dict[str, Any] = {}
        total_rows = 0

        target_sheets = [sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.sheetnames
        for sname in target_sheets:
            ws = wb[sname]
            rows: list[list[Any]] = []
            for row in ws.iter_rows(max_row=max_rows, values_only=True):
                rows.append([cell if cell is not None else "" for cell in row])
                total_rows += 1
            sheets_data[sname] = {"rows": rows, "row_count": len(rows)}

        wb.close()
        return {
            "success": True,
            "data": {
                "file_path": file_path,
                "all_sheets": wb.sheetnames,
                "sheets_read": list(sheets_data.keys()),
                "sheets_data": sheets_data,
                "total_rows": total_rows,
            },
        }

    # --- Create tools (synchronous, executed in thread pool) ---

    def _create_pdf(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _FPDF_AVAILABLE:
            return {"success": False, "error": "fpdf2 not installed. Run: pip install fpdf2"}

        from fpdf import FPDF

        filename: str = args["filename"]
        title: str = args.get("title", "")
        content: str = args["content"]
        output_dir: str = args.get("output_dir", _DEFAULT_OUTPUT_DIR)

        out_path = _ensure_output_dir(output_dir) / filename

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Use built-in font for basic Latin; for non-Latin (Vietnamese diacritics)
        # callers should supply pre-encoded text or install a Unicode font.
        # epw = effective page width (page width minus left and right margins)
        if title:
            pdf.set_font("Helvetica", size=16, style="B")
            pdf.multi_cell(pdf.epw, 10, title, align="C")
            pdf.ln(5)

        pdf.set_font("Helvetica", size=11)
        for line in content.split("\n"):
            pdf.multi_cell(pdf.epw, 8, line if line else " ")

        pdf.output(str(out_path))
        return {
            "success": True,
            "data": {
                "file_path": str(out_path),
                "filename": filename,
                "size_bytes": out_path.stat().st_size,
            },
        }

    def _create_word(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        import docx as _docx

        filename: str = args["filename"]
        title: str = args.get("title", "")
        sections: list[dict[str, str]] = args.get("sections", [])
        output_dir: str = args.get("output_dir", _DEFAULT_OUTPUT_DIR)

        out_path = _ensure_output_dir(output_dir) / filename

        doc = _docx.Document()
        if title:
            doc.add_heading(title, level=1)

        for section in sections:
            heading = section.get("heading", "")
            body = section.get("body", section.get("text", ""))
            if heading:
                doc.add_heading(heading, level=2)
            if body:
                doc.add_paragraph(body)

        doc.save(str(out_path))
        return {
            "success": True,
            "data": {
                "file_path": str(out_path),
                "filename": filename,
                "size_bytes": out_path.stat().st_size,
            },
        }

    def _create_excel(self, args: dict[str, Any]) -> dict[str, Any]:
        if not _OPENPYXL_AVAILABLE:
            return {"success": False, "error": "openpyxl not installed. Run: pip install openpyxl"}

        import openpyxl as _openpyxl

        filename: str = args["filename"]
        sheet_name: str = args.get("sheet_name", "Sheet1")
        headers: list[str] = args.get("headers", [])
        rows: list[list[Any]] = args.get("rows", [])
        output_dir: str = args.get("output_dir", _DEFAULT_OUTPUT_DIR)

        out_path = _ensure_output_dir(output_dir) / filename

        wb = _openpyxl.Workbook()
        ws = wb.active
        if ws is None:
            ws = wb.create_sheet()
        ws.title = sheet_name

        if headers:
            ws.append(headers)
        for row in rows:
            ws.append(row)

        wb.save(str(out_path))
        return {
            "success": True,
            "data": {
                "file_path": str(out_path),
                "filename": filename,
                "sheet_name": sheet_name,
                "row_count": len(rows),
                "size_bytes": out_path.stat().st_size,
            },
        }

    # --- Summarize (async, chains read + LLM) ---

    async def _summarize_document(self, args: dict[str, Any]) -> dict[str, Any]:
        file_path: str = args["file_path"]
        language: str = args.get("language", "")
        focus: str = args.get("focus", "")

        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            read_result = await asyncio.get_event_loop().run_in_executor(
                None, self._read_pdf, {"file_path": file_path}
            )
        elif suffix in (".docx", ".doc"):
            read_result = await asyncio.get_event_loop().run_in_executor(
                None, self._read_word, {"file_path": file_path}
            )
        elif suffix in (".xlsx", ".xls", ".xlsm"):
            read_result = await asyncio.get_event_loop().run_in_executor(
                None, self._read_excel, {"file_path": file_path}
            )
        else:
            return {
                "success": False,
                "error": f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, .xlsx/.xls",
            }

        if not read_result.get("success"):
            return read_result

        raw_text = read_result["data"].get("text", "")
        if not raw_text.strip():
            return {
                "success": False,
                "error": "No text could be extracted from the document.",
            }

        if not self._llm_router:
            snippet = raw_text[:500]
            return {
                "success": True,
                "data": {
                    "file_path": file_path,
                    "summary": f"[LLM unavailable] First 500 chars:\n{snippet}",
                    "llm_available": False,
                },
            }

        lang_instruction = f" Respond in {language}." if language else ""
        focus_instruction = f" Focus especially on: {focus}." if focus else ""
        system_prompt = (
            "You are an expert document analyst. Produce a clear, concise summary "
            "of the document provided. Highlight the key points, main conclusions, "
            f"and any important data.{lang_instruction}{focus_instruction}"
        )

        llm_resp = await self.call_llm(
            LLMRequest(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"Document ({Path(file_path).name}):\n\n{raw_text[:10_000]}",
                    },
                ],
                temperature=0.2,
                max_tokens=1024,
            )
        )

        return {
            "success": True,
            "data": {
                "file_path": file_path,
                "file_type": suffix,
                "summary": llm_resp.content,
                "char_count": len(raw_text),
                "llm_model": llm_resp.model,
            },
        }

    # ------------------------------------------------------------------
    # LLM-unavailable fallback
    # ------------------------------------------------------------------

    def _no_llm_response(self, request: AgentRequest, error: str) -> AgentResponse:
        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=False,
            error=(
                "LLM is unavailable for intent routing. "
                "Pass 'tool' and 'args' directly in input_data, or ensure an LLM is configured. "
                f"LLM error: {error}"
            ),
        )
