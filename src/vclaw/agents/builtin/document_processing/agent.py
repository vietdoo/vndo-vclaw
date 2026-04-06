"""Document processing agent for reading and creating PDF, DOCX, and XLSX files.

Library selection (graceful degradation):
  - pypdf           — PDF text extraction
  - python-docx     — DOCX read/write
  - openpyxl        — XLSX read/write
  - reportlab       — PDF creation (optional, falls back to text-only)

Install all backends:
    pip install pypdf python-docx openpyxl reportlab
"""

from __future__ import annotations

import base64
import io
import json
import os
import tempfile
from pathlib import Path
from typing import Any, ClassVar

import structlog

from vclaw.agents.base import AgentBase
from vclaw.domain.models import (
    AgentCapability,
    AgentManifest,
    AgentRequest,
    AgentResponse,
    LLMRequest,
    RetryPolicy,
    ToolDefinition,
)

logger: structlog.stdlib.BoundLogger = structlog.get_logger(__name__)

try:
    import pypdf

    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

try:
    import docx as python_docx

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import openpyxl

    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as rl_canvas

    _REPORTLAB_AVAILABLE = True
except ImportError:
    _REPORTLAB_AVAILABLE = False


_OUTPUT_DIR = os.environ.get("VCLAW_DOCUMENT_OUTPUT_DIR", tempfile.gettempdir())


class DocumentProcessingAgent(AgentBase):
    """Agent for reading and creating document files (PDF, DOCX, XLSX).

    Reads text content from uploaded documents and creates simple documents
    from structured data. Supports base64-encoded file input and produces
    base64-encoded file output for transport over the event bus.
    """

    manifest: ClassVar[AgentManifest] = AgentManifest(
        name="document_processing",
        version="0.1.0",
        description=(
            "Reads and creates document files: extract text from PDF, DOCX, XLSX; "
            "create new PDF, DOCX, XLSX files from structured data. "
            "Supports Vietnamese and English content."
        ),
        capabilities=[
            AgentCapability(
                name="document_reading",
                description=(
                    "Extract text content from PDF, DOCX, and XLSX files. "
                    "Reads uploaded documents and returns structured text."
                ),
            ),
            AgentCapability(
                name="document_creation",
                description=("Create new PDF, DOCX, or XLSX files from provided text, tables, or structured data."),
            ),
            AgentCapability(
                name="document_info",
                description="Get metadata and summary information about a document file.",
            ),
        ],
        tools=[
            ToolDefinition(
                name="read_pdf",
                description="Extract text content from a PDF file (base64-encoded or file path)",
                parameters={
                    "file_base64": {
                        "type": "string",
                        "description": "Base64-encoded PDF file content",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to PDF file on disk",
                    },
                    "page_numbers": {
                        "type": "array",
                        "items": {"type": "integer"},
                        "description": "Specific page numbers to extract (0-indexed). Empty = all pages.",
                    },
                },
                required_params=[],
            ),
            ToolDefinition(
                name="read_docx",
                description="Extract text content from a DOCX (Word) file",
                parameters={
                    "file_base64": {
                        "type": "string",
                        "description": "Base64-encoded DOCX file content",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to DOCX file on disk",
                    },
                    "include_tables": {
                        "type": "boolean",
                        "description": "Whether to extract table data (default: true)",
                    },
                },
                required_params=[],
            ),
            ToolDefinition(
                name="read_xlsx",
                description="Extract data from an XLSX (Excel) spreadsheet",
                parameters={
                    "file_base64": {
                        "type": "string",
                        "description": "Base64-encoded XLSX file content",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to XLSX file on disk",
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "Specific sheet name to read. Empty = first sheet.",
                    },
                    "max_rows": {
                        "type": "integer",
                        "description": "Maximum rows to read (default: 500)",
                    },
                },
                required_params=[],
            ),
            ToolDefinition(
                name="create_pdf",
                description="Create a new PDF file from text content",
                parameters={
                    "title": {"type": "string", "description": "Document title"},
                    "content": {
                        "type": "string",
                        "description": "Text content for the PDF (supports newlines for paragraphs)",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: output.pdf)",
                    },
                },
                required_params=["content"],
            ),
            ToolDefinition(
                name="create_docx",
                description="Create a new DOCX (Word) document",
                parameters={
                    "title": {"type": "string", "description": "Document title (added as heading)"},
                    "content": {
                        "type": "string",
                        "description": "Text content (paragraphs separated by newlines)",
                    },
                    "sections": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "heading": {"type": "string"},
                                "body": {"type": "string"},
                            },
                        },
                        "description": "Structured sections with headings and body text",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: output.docx)",
                    },
                },
                required_params=["content"],
            ),
            ToolDefinition(
                name="create_xlsx",
                description="Create a new XLSX (Excel) spreadsheet",
                parameters={
                    "title": {"type": "string", "description": "Sheet title / name"},
                    "headers": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Column headers",
                    },
                    "rows": {
                        "type": "array",
                        "items": {"type": "array"},
                        "description": "Data rows (array of arrays)",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: output.xlsx)",
                    },
                },
                required_params=["headers", "rows"],
            ),
            ToolDefinition(
                name="get_document_info",
                description="Get metadata about a document (page count, sheet names, word count, etc.)",
                parameters={
                    "file_base64": {
                        "type": "string",
                        "description": "Base64-encoded file content",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to the file on disk",
                    },
                    "file_type": {
                        "type": "string",
                        "enum": ["pdf", "docx", "xlsx"],
                        "description": "File format",
                    },
                },
                required_params=["file_type"],
            ),
        ],
        max_concurrent=5,
        timeout_seconds=120.0,
        retry_policy=RetryPolicy(max_retries=2, base_delay_seconds=1.0),
        tags=["document", "pdf", "docx", "xlsx", "file_processing"],
    )

    async def execute(self, request: AgentRequest) -> AgentResponse:
        """Process document request via LLM tool calling or direct dispatch."""
        text = request.input_data.get("text", "")
        file_base64 = request.input_data.get("file_base64", "")
        file_path = request.input_data.get("file_path", "")
        file_type = request.input_data.get("file_type", "")

        if not text and not file_base64 and not file_path:
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=False,
                error="No input text or file provided",
            )

        if file_base64 or file_path:
            return await self._direct_file_processing(request, file_base64, file_path, file_type, text)

        try:
            llm_resp = await self.call_llm(
                LLMRequest(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a document processing assistant. Use the available tools to "
                                "read or create documents (PDF, DOCX, XLSX). Always use a tool call.\n"
                                f"Available libraries: PDF read={_PYPDF_AVAILABLE}, "
                                f"DOCX={_DOCX_AVAILABLE}, XLSX={_OPENPYXL_AVAILABLE}, "
                                f"PDF create={_REPORTLAB_AVAILABLE}"
                            ),
                        },
                        {"role": "user", "content": text},
                    ],
                    tools=self.get_tool_schemas(),
                    tool_choice="auto",
                    temperature=0.0,
                )
            )
        except Exception as exc:
            return await self._fallback_execution(request, str(exc))

        if llm_resp.tool_calls:
            return await self._handle_tool_calls(request, llm_resp.tool_calls)

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=True,
            data={"response_text": llm_resp.content or "Document operation completed."},
        )

    async def _direct_file_processing(
        self,
        request: AgentRequest,
        file_base64: str,
        file_path: str,
        file_type: str,
        text: str,
    ) -> AgentResponse:
        """Direct file read when a file is provided in input_data."""
        if not file_type:
            file_type = self._detect_file_type(file_path)

        tool_map = {"pdf": "read_pdf", "docx": "read_docx", "xlsx": "read_xlsx"}
        tool_name = tool_map.get(file_type, "")

        if not tool_name:
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=False,
                error=f"Unsupported file type: '{file_type}'. Supported: pdf, docx, xlsx",
            )

        args: dict[str, Any] = {}
        if file_base64:
            args["file_base64"] = file_base64
        if file_path:
            args["file_path"] = file_path

        result = await self._execute_tool(tool_name, args)
        if result.get("success"):
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=True,
                data={
                    "response_text": self._format_tool_result(tool_name, result["data"]),
                    "tool_results": [{"tool": tool_name, "result": result}],
                },
            )

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=False,
            error=result.get("error", "Failed to process document"),
        )

    async def _handle_tool_calls(self, request: AgentRequest, tool_calls: list[dict[str, Any]]) -> AgentResponse:
        results: list[dict[str, Any]] = []
        response_parts: list[str] = []

        for tc in tool_calls:
            func = tc.get("function", {})
            name = func.get("name", "")
            try:
                args: dict[str, Any] = json.loads(func.get("arguments", "{}"))
            except json.JSONDecodeError:
                args = {}

            result = await self._execute_tool(name, args)
            results.append({"tool": name, "result": result})

            if result.get("success"):
                response_parts.append(self._format_tool_result(name, result["data"]))
            else:
                response_parts.append(f"❌ {result.get('error', 'Unknown error')}")

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=True,
            data={
                "response_text": "\n\n".join(response_parts),
                "tool_results": results,
            },
        )

    def _format_tool_result(self, tool_name: str, data: Any) -> str:
        if not isinstance(data, dict):
            return str(data)

        if tool_name in ("read_pdf", "read_docx"):
            pages = data.get("page_count") or data.get("paragraph_count", "?")
            text = data.get("text", "")[:2000]
            label = "📄 PDF" if tool_name == "read_pdf" else "📝 DOCX"
            info = f"pages: {pages}" if tool_name == "read_pdf" else f"paragraphs: {pages}"
            parts = [f"{label} ({info})", "---", text]
            tables = data.get("tables")
            if tables:
                parts.append(f"\n📊 Tables found: {len(tables)}")
            return "\n".join(parts)

        if tool_name == "read_xlsx":
            sheet = data.get("sheet_name", "?")
            row_count = data.get("row_count", 0)
            headers = data.get("headers", [])
            preview = data.get("preview_rows", [])[:5]
            parts = [f"📊 Excel sheet '{sheet}' ({row_count} rows)", f"Headers: {', '.join(str(h) for h in headers)}"]
            for row in preview:
                parts.append(f"  | {' | '.join(str(c) for c in row)} |")
            return "\n".join(parts)

        if tool_name in ("create_pdf", "create_docx", "create_xlsx"):
            fname = data.get("filename", "?")
            size = data.get("size_bytes", 0)
            return f"✅ Created: {fname} ({size:,} bytes)"

        if tool_name == "get_document_info":
            parts = [f"📋 Document info ({data.get('file_type', '?')})"]
            for k, v in data.items():
                if k not in ("file_type",):
                    parts.append(f"  • {k}: {v}")
            return "\n".join(parts)

        return str(data)

    async def _execute_tool(self, name: str, args: dict[str, Any]) -> dict[str, Any]:
        try:
            match name:
                case "read_pdf":
                    return self._read_pdf(**args)
                case "read_docx":
                    return self._read_docx(**args)
                case "read_xlsx":
                    return self._read_xlsx(**args)
                case "create_pdf":
                    return self._create_pdf(**args)
                case "create_docx":
                    return self._create_docx(**args)
                case "create_xlsx":
                    return self._create_xlsx(**args)
                case "get_document_info":
                    return self._get_document_info(**args)
                case _:
                    return {"success": False, "error": f"Unknown tool: {name}"}
        except Exception as exc:
            logger.exception("document_tool_error", tool=name)
            return {"success": False, "error": str(exc)}

    def _resolve_file(self, file_base64: str = "", file_path: str = "") -> io.BytesIO:
        """Resolve a file from base64 content or disk path into a BytesIO stream."""
        if file_base64:
            return io.BytesIO(base64.b64decode(file_base64))
        if file_path:
            path = Path(file_path)
            if not path.is_file():
                raise FileNotFoundError(f"File not found: {file_path}")
            return io.BytesIO(path.read_bytes())
        raise ValueError("Either file_base64 or file_path must be provided")

    @staticmethod
    def _detect_file_type(file_path: str) -> str:
        if not file_path:
            return ""
        ext = Path(file_path).suffix.lower().lstrip(".")
        return {"doc": "docx", "xls": "xlsx"}.get(ext, ext)

    def _read_pdf(
        self,
        file_base64: str = "",
        file_path: str = "",
        page_numbers: list[int] | None = None,
        **_: Any,
    ) -> dict[str, Any]:
        if not _PYPDF_AVAILABLE:
            return {"success": False, "error": "pypdf not installed. Run: pip install pypdf"}

        stream = self._resolve_file(file_base64, file_path)
        reader = pypdf.PdfReader(stream)
        total_pages = len(reader.pages)

        pages_to_read = page_numbers if page_numbers else list(range(total_pages))
        text_parts: list[str] = []

        for page_num in pages_to_read:
            if 0 <= page_num < total_pages:
                page_text = reader.pages[page_num].extract_text() or ""
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

        metadata = reader.metadata
        meta_dict: dict[str, Any] = {}
        if metadata:
            for key in ("/Title", "/Author", "/Subject", "/Creator"):
                val = metadata.get(key)
                if val:
                    meta_dict[key.strip("/")] = str(val)

        return {
            "success": True,
            "data": {
                "text": "\n\n".join(text_parts),
                "page_count": total_pages,
                "pages_read": len(text_parts),
                "metadata": meta_dict,
                "word_count": sum(len(p.split()) for p in text_parts),
            },
        }

    def _read_docx(
        self,
        file_base64: str = "",
        file_path: str = "",
        include_tables: bool = True,
        **_: Any,
    ) -> dict[str, Any]:
        if not _DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        stream = self._resolve_file(file_base64, file_path)
        doc = python_docx.Document(stream)

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name:
                    paragraphs.append(f"## {text}")
                else:
                    paragraphs.append(text)

        tables_data: list[list[list[str]]] = []
        if include_tables:
            for table in doc.tables:
                rows: list[list[str]] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                tables_data.append(rows)

        core = doc.core_properties
        meta: dict[str, Any] = {}
        if core.title:
            meta["title"] = core.title
        if core.author:
            meta["author"] = core.author

        return {
            "success": True,
            "data": {
                "text": "\n\n".join(paragraphs),
                "paragraph_count": len(paragraphs),
                "tables": tables_data,
                "table_count": len(tables_data),
                "metadata": meta,
                "word_count": sum(len(p.split()) for p in paragraphs),
            },
        }

    def _read_xlsx(
        self,
        file_base64: str = "",
        file_path: str = "",
        sheet_name: str = "",
        max_rows: int = 500,
        **_: Any,
    ) -> dict[str, Any]:
        if not _OPENPYXL_AVAILABLE:
            return {"success": False, "error": "openpyxl not installed. Run: pip install openpyxl"}

        stream = self._resolve_file(file_base64, file_path)
        wb = openpyxl.load_workbook(stream, read_only=True, data_only=True)

        target_sheet = sheet_name if sheet_name and sheet_name in wb.sheetnames else wb.sheetnames[0]
        ws = wb[target_sheet]

        all_rows: list[list[Any]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows + 1:
                break
            all_rows.append([cell if cell is not None else "" for cell in row])

        wb.close()

        headers = [str(h) for h in all_rows[0]] if all_rows else []
        data_rows = all_rows[1:] if len(all_rows) > 1 else []

        return {
            "success": True,
            "data": {
                "sheet_name": target_sheet,
                "sheet_names": wb.sheetnames if hasattr(wb, "sheetnames") else [target_sheet],
                "headers": headers,
                "rows": [[str(c) for c in row] for row in data_rows],
                "row_count": len(data_rows),
                "column_count": len(headers),
                "preview_rows": [[str(c) for c in row] for row in data_rows[:10]],
            },
        }

    def _create_pdf(
        self,
        content: str,
        title: str = "",
        filename: str = "output.pdf",
        **_: Any,
    ) -> dict[str, Any]:
        if not _REPORTLAB_AVAILABLE:
            return {"success": False, "error": "reportlab not installed. Run: pip install reportlab"}

        output_path = Path(_OUTPUT_DIR) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)

        c = rl_canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4

        y_pos = height - 40 * mm
        left_margin = 25 * mm
        usable_width = width - 2 * left_margin

        if title:
            c.setFont("Helvetica-Bold", 16)
            c.drawString(left_margin, y_pos, title)
            y_pos -= 12 * mm

        c.setFont("Helvetica", 11)
        line_height = 5 * mm

        for paragraph in content.split("\n"):
            words = paragraph.split()
            if not words:
                y_pos -= line_height
                if y_pos < 30 * mm:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y_pos = height - 25 * mm
                continue

            line = ""
            for word in words:
                test_line = f"{line} {word}".strip()
                if c.stringWidth(test_line, "Helvetica", 11) > usable_width:
                    c.drawString(left_margin, y_pos, line)
                    y_pos -= line_height
                    line = word
                    if y_pos < 30 * mm:
                        c.showPage()
                        c.setFont("Helvetica", 11)
                        y_pos = height - 25 * mm
                else:
                    line = test_line

            if line:
                c.drawString(left_margin, y_pos, line)
                y_pos -= line_height
                if y_pos < 30 * mm:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y_pos = height - 25 * mm

        c.save()

        file_bytes = output_path.read_bytes()
        return {
            "success": True,
            "data": {
                "filename": filename,
                "file_path": str(output_path),
                "size_bytes": len(file_bytes),
                "file_base64": base64.b64encode(file_bytes).decode("ascii"),
            },
        }

    def _create_docx(
        self,
        content: str,
        title: str = "",
        sections: list[dict[str, str]] | None = None,
        filename: str = "output.docx",
        **_: Any,
    ) -> dict[str, Any]:
        if not _DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        doc = python_docx.Document()

        if title:
            doc.add_heading(title, level=0)

        if sections:
            for section in sections:
                heading = section.get("heading", "")
                body = section.get("body", "")
                if heading:
                    doc.add_heading(heading, level=1)
                if body:
                    for para in body.split("\n"):
                        if para.strip():
                            doc.add_paragraph(para.strip())
        else:
            for para in content.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        output_path = Path(_OUTPUT_DIR) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        file_bytes = output_path.read_bytes()
        return {
            "success": True,
            "data": {
                "filename": filename,
                "file_path": str(output_path),
                "size_bytes": len(file_bytes),
                "file_base64": base64.b64encode(file_bytes).decode("ascii"),
            },
        }

    def _create_xlsx(
        self,
        headers: list[str],
        rows: list[list[Any]],
        title: str = "Sheet1",
        filename: str = "output.xlsx",
        **_: Any,
    ) -> dict[str, Any]:
        if not _OPENPYXL_AVAILABLE:
            return {"success": False, "error": "openpyxl not installed. Run: pip install openpyxl"}

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title or "Sheet1"

        ws.append(headers)

        from openpyxl.styles import Font

        for cell in ws[1]:
            cell.font = Font(bold=True)

        for row in rows:
            ws.append(row)

        for col_cells in ws.columns:
            max_len = 0
            for cell in col_cells:
                val = str(cell.value) if cell.value else ""
                max_len = max(max_len, len(val))
            adjusted = min(max_len + 2, 50)
            ws.column_dimensions[col_cells[0].column_letter].width = adjusted

        output_path = Path(_OUTPUT_DIR) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(output_path))

        file_bytes = output_path.read_bytes()
        return {
            "success": True,
            "data": {
                "filename": filename,
                "file_path": str(output_path),
                "size_bytes": len(file_bytes),
                "file_base64": base64.b64encode(file_bytes).decode("ascii"),
            },
        }

    def _get_document_info(
        self,
        file_type: str,
        file_base64: str = "",
        file_path: str = "",
        **_: Any,
    ) -> dict[str, Any]:
        stream = self._resolve_file(file_base64, file_path)

        match file_type.lower():
            case "pdf":
                if not _PYPDF_AVAILABLE:
                    return {"success": False, "error": "pypdf not installed"}
                reader = pypdf.PdfReader(stream)
                meta = reader.metadata or {}
                return {
                    "success": True,
                    "data": {
                        "file_type": "pdf",
                        "page_count": len(reader.pages),
                        "title": str(meta.get("/Title", "")),
                        "author": str(meta.get("/Author", "")),
                        "encrypted": reader.is_encrypted,
                    },
                }
            case "docx":
                if not _DOCX_AVAILABLE:
                    return {"success": False, "error": "python-docx not installed"}
                doc = python_docx.Document(stream)
                return {
                    "success": True,
                    "data": {
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "title": doc.core_properties.title or "",
                        "author": doc.core_properties.author or "",
                    },
                }
            case "xlsx":
                if not _OPENPYXL_AVAILABLE:
                    return {"success": False, "error": "openpyxl not installed"}
                wb = openpyxl.load_workbook(stream, read_only=True)
                sheets_info = {}
                for name in wb.sheetnames:
                    ws = wb[name]
                    sheets_info[name] = {
                        "max_row": ws.max_row,
                        "max_column": ws.max_column,
                    }
                wb.close()
                return {
                    "success": True,
                    "data": {
                        "file_type": "xlsx",
                        "sheet_count": len(wb.sheetnames),
                        "sheets": sheets_info,
                    },
                }
            case _:
                return {"success": False, "error": f"Unsupported file type: {file_type}"}

    async def _fallback_execution(self, request: AgentRequest, error: str) -> AgentResponse:
        """Keyword-based fallback when LLM is unavailable."""
        text = request.input_data.get("text", "").lower()

        if any(kw in text for kw in ["đọc", "read", "extract", "trích xuất"]):
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=False,
                error=(
                    "To read a document, please provide the file via file_base64 or file_path "
                    f"in input_data. LLM unavailable: {error}"
                ),
            )

        if any(kw in text for kw in ["tạo", "create", "viết", "write"]):
            return AgentResponse(
                workflow_id=request.workflow_id,
                subtask_id=request.subtask_id,
                agent_name=self.name,
                success=False,
                error=(f"LLM is needed to parse document creation requests. LLM error: {error}"),
            )

        return AgentResponse(
            workflow_id=request.workflow_id,
            subtask_id=request.subtask_id,
            agent_name=self.name,
            success=False,
            error=(f"Supported operations: read PDF/DOCX/XLSX, create PDF/DOCX/XLSX. LLM unavailable: {error}"),
        )
